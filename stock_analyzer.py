"""
stock_analyzer.py
=================
Minervini + Kell + POC + PPV + EP + ATR technical analysis -> concise DOCX report.

Report format: actionable narrative (2 paragraphs) — bullish picture,
bearish picture, and a recommended action with reclaim levels.
NO scorecards, checkboxes, or raw numbers in the output — the computer
does the math, you get the decision.

USAGE:
    pip install ib_insync yfinance pandas numpy scipy python-docx anthropic
    python stock_analyzer.py AAPL
    python stock_analyzer.py HOOD --chart HOOD.png        # includes Claude vision
    python stock_analyzer.py SE --period 1y --output SE.docx
"""

import argparse
import sys
import os
import io
from datetime import datetime, timedelta

if sys.platform == "win32":
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")
    except Exception:
        pass
    # Enable ANSI escape sequences (colors) in Windows terminals
    try:
        import ctypes
        kernel32 = ctypes.windll.kernel32
        kernel32.SetConsoleMode(kernel32.GetStdHandle(-11), 7)
    except Exception:
        pass

import numpy as np
import pandas as pd

try:
    import yfinance as yf
    HAS_YFINANCE = True
except ImportError:
    HAS_YFINANCE = False

try:
    from ib_insync import IB, Stock, util
    HAS_IBKR = True
except ImportError:
    HAS_IBKR = False

try:
    from scipy.signal import find_peaks
except ImportError:
    print("ERROR: scipy not installed. Run: pip install scipy")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# ============================================================================
# DATA FETCH
# ============================================================================

def fetch_data_tws(ticker, period="3y", host="127.0.0.1", port=7497, client_id=20):
    import asyncio
    print(f"[fetch] TWS connecting {host}:{port} clientId={client_id}...")
    try:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        util.patchAsyncio()
    except Exception:
        pass

    ib = IB()
    try:
        ib.connect(host, port, clientId=client_id, timeout=15)
    except Exception as e:
        raise RuntimeError(f"TWS connection failed: {e}")

    try:
        contract = Stock(ticker, "SMART", "USD")
        details = ib.reqContractDetails(contract)
        if not details:
            for primary in ("NYSE", "NASDAQ", "ARCA", "BATS"):
                alt = Stock(ticker, "SMART", "USD", primaryExchange=primary)
                details = ib.reqContractDetails(alt)
                if details:
                    contract = alt
                    break
        if not details:
            raise RuntimeError(f"No contract details for {ticker}")

        contract = details[0].contract
        print(f"[fetch] TWS qualified {ticker}: {contract.primaryExchange or contract.exchange}")

        duration_map = {"1y": "1 Y", "2y": "2 Y", "3y": "3 Y", "5y": "5 Y", "10y": "10 Y", "max": "15 Y"}
        duration = duration_map.get(period.lower(), "3 Y")

        bars = ib.reqHistoricalData(
            contract, endDateTime="", durationStr=duration,
            barSizeSetting="1 day", whatToShow="TRADES", useRTH=True, formatDate=1,
        )
        if not bars:
            raise RuntimeError(f"TWS returned no bars for {ticker}")

        df = util.df(bars)
        df = df.rename(columns={"date": "Date", "open": "Open", "high": "High",
                                "low": "Low", "close": "Close", "volume": "Volume"})
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.set_index("Date")[["Open", "High", "Low", "Close", "Volume"]]
        if df.index.tz is not None:
            df.index = df.index.tz_localize(None)
        print(f"[fetch] Got {len(df)} bars from TWS, {df.index[0].date()} to {df.index[-1].date()}")
        return df
    finally:
        try:
            ib.disconnect()
        except Exception:
            pass


def fetch_data_yfinance(ticker, period="3y", retries=3):
    import time
    print(f"[fetch] yfinance downloading {ticker} - period={period}...")
    t = yf.Ticker(ticker)
    last_err = None
    df = pd.DataFrame()
    for attempt in range(retries):
        try:
            df = t.history(period=period, interval="1d", auto_adjust=False)
            if not df.empty:
                break
        except Exception as e:
            last_err = e
            if "rate" in str(e).lower() or "too many" in str(e).lower():
                wait = 30 * (attempt + 1)
                print(f"[fetch] Rate limited, waiting {wait}s...")
                time.sleep(wait)
            else:
                raise
    if df.empty:
        raise RuntimeError(f"Failed to fetch {ticker} via yfinance. Last error: {last_err}")
    if isinstance(df.columns, pd.MultiIndex):
        df.columns = df.columns.get_level_values(0)
    df = df[["Open", "High", "Low", "Close", "Volume"]].copy()
    df.index = pd.to_datetime(df.index)
    if df.index.tz is not None:
        df.index = df.index.tz_localize(None)
    print(f"[fetch] Got {len(df)} bars from yfinance, {df.index[0].date()} to {df.index[-1].date()}")
    return df


def fetch_data(ticker, period="3y", source="tws", tws_host="127.0.0.1",
               tws_port=7497, tws_client=20):
    df = None
    yf_ticker = None
    if source in ("tws", "tws_only") and HAS_IBKR:
        try:
            df = fetch_data_tws(ticker, period, tws_host, tws_port, tws_client)
        except Exception as e:
            print(f"[fetch] TWS failed: {e}")
            if source == "tws_only":
                raise
            print("[fetch] Falling back to yfinance...")
    if df is None:
        if not HAS_YFINANCE:
            raise RuntimeError("TWS unavailable and yfinance not installed.")
        df = fetch_data_yfinance(ticker, period)
    if HAS_YFINANCE:
        try:
            yf_ticker = yf.Ticker(ticker)
        except Exception:
            yf_ticker = None
    return df, yf_ticker


def fetch_news(ticker_obj, days=45):
    try:
        news = ticker_obj.news or []
    except Exception:
        return []
    cutoff = datetime.now() - timedelta(days=days)
    filtered = []
    for item in news:
        content = item.get("content", item)
        title = content.get("title", "")
        pub_raw = content.get("pubDate") or content.get("providerPublishTime")
        if pub_raw:
            try:
                if isinstance(pub_raw, (int, float)):
                    pub = datetime.fromtimestamp(pub_raw)
                else:
                    pub = pd.to_datetime(pub_raw).to_pydatetime()
                    if pub.tzinfo is not None:
                        pub = pub.replace(tzinfo=None)
            except Exception:
                continue
            if pub >= cutoff and title:
                publisher = (content.get("provider", {}).get("displayName")
                             or content.get("publisher", "") or "Unknown")
                filtered.append({"date": pub, "title": title, "publisher": publisher})
    filtered.sort(key=lambda x: x["date"], reverse=True)
    return filtered


# ============================================================================
# INDICATORS
# ============================================================================

def _atr(df, period=14):
    hl = df["High"] - df["Low"]
    hc = (df["High"] - df["Close"].shift()).abs()
    lc = (df["Low"] - df["Close"].shift()).abs()
    tr = pd.concat([hl, hc, lc], axis=1).max(axis=1)
    return tr.rolling(period).mean()


def add_indicators(df):
    df = df.copy()
    df["EMA8"] = df["Close"].ewm(span=8, adjust=False).mean()
    df["EMA21"] = df["Close"].ewm(span=21, adjust=False).mean()
    df["SMA50"] = df["Close"].rolling(50).mean()
    df["SMA200"] = df["Close"].rolling(200).mean()
    df["SMA50_slope20"] = df["SMA50"].diff(20)
    df["SMA200_slope20"] = df["SMA200"].diff(20)
    df["ATR14"] = _atr(df, 14)
    df["VolSMA50"] = df["Volume"].rolling(50).mean()
    df["RVol"] = df["Volume"] / df["VolSMA50"]
    df["PctChg"] = df["Close"].pct_change() * 100
    df["UpDay"] = df["Close"] > df["Open"]
    df["ATRMult"] = (df["Close"] - df["SMA50"]) / df["ATR14"]
    return df


def classify_volume_colors(df):
    df = df.copy()
    colors = []
    for i in range(len(df)):
        if i < 10:
            colors.append("gray")
            continue
        lookback = df.iloc[max(0, i - 10):i]
        down_days = lookback[lookback["Close"] < lookback["Open"]]
        max_down = down_days["Volume"].max() if not down_days.empty else 0
        row = df.iloc[i]
        vol, vol_sma = row["Volume"], row["VolSMA50"]
        up = row["Close"] > row["Open"]
        if pd.isna(vol_sma) or vol_sma == 0:
            colors.append("gray")
            continue
        if up and vol > max_down and vol > vol_sma:
            colors.append("blue")
        elif up and vol > vol_sma:
            colors.append("green")
        elif not up and vol > vol_sma:
            colors.append("red")
        elif vol < vol_sma * 0.2:
            colors.append("orange")
        else:
            colors.append("gray")
    df["VolColor"] = colors

    df["HVQ"] = False
    df["HVY"] = False
    for i in range(63, len(df)):
        row = df.iloc[i]
        if not row["UpDay"]:
            continue
        if row["Volume"] == df["Volume"].iloc[max(0, i - 63):i + 1].max():
            df.loc[df.index[i], "HVQ"] = True
        if i >= 252 and row["Volume"] == df["Volume"].iloc[max(0, i - 252):i + 1].max():
            df.loc[df.index[i], "HVY"] = True
    return df


# ============================================================================
# STRUCTURE / SWINGS / RECLAIM TIERS
# ============================================================================

def find_swings(df, prominence_pct=0.05, distance=5):
    highs = df["High"].values
    lows = df["Low"].values
    prom = df["Close"].mean() * prominence_pct
    hi_idx, _ = find_peaks(highs, prominence=prom, distance=distance)
    lo_idx, _ = find_peaks(-lows, prominence=prom, distance=distance)
    return list(hi_idx), list(lo_idx)


def determine_structure(df, swing_highs, swing_lows):
    if len(swing_highs) < 2 or len(swing_lows) < 2:
        return "insufficient swings", None, None
    hh1 = df.iloc[swing_highs[-2]]["High"]
    hh2 = df.iloc[swing_highs[-1]]["High"]
    ll1 = df.iloc[swing_lows[-2]]["Low"]
    ll2 = df.iloc[swing_lows[-1]]["Low"]
    prior_hi, prior_lo = float(hh2), float(ll2)
    price = df["Close"].iloc[-1]
    if hh2 > hh1 and ll2 > ll1:
        s = "Uptrend (higher highs + higher lows)"
    elif hh2 < hh1 and ll2 < ll1:
        s = "Downtrend (lower highs + lower lows)"
    elif price > prior_hi:
        s = "Reclaimed prior swing high — downtrend broken"
    else:
        s = "Mixed / transitioning"
    return s, prior_hi, prior_lo


def determine_reclaim_levels(df, swing_highs):
    tiers = {}
    tiers["tier1"] = {"price": float(df["High"].tail(10).max()),
                      "desc": "Last 10-bar high (short-term swing)"}
    lb_start = max(0, len(df) - 60)
    lb_end = max(0, len(df) - 10)
    if lb_end > lb_start:
        tiers["tier2"] = {"price": float(df.iloc[lb_start:lb_end]["High"].max()),
                          "desc": "3-month high prior to current consolidation"}
    else:
        tiers["tier2"] = {"price": tiers["tier1"]["price"], "desc": "Same as Tier 1"}
    lb6 = df.tail(126)
    down = lb6[lb6["Close"] < lb6["Open"]]
    if not down.empty:
        worst = down["Volume"].idxmax()
        pos = df.index.get_loc(worst)
        if pos > 0:
            prior = df.iloc[pos - 1]
            tiers["tier3"] = {"price": float(prior["High"]),
                              "desc": f"High of bar before largest distribution day ({worst.date()})"}
        else:
            tiers["tier3"] = {"price": float(lb6["High"].max()), "desc": "6-month high"}
    else:
        tiers["tier3"] = {"price": float(lb6["High"].max()), "desc": "6-month high"}
    return tiers


# ============================================================================
# STAGE
# ============================================================================

def determine_stage(df):
    last = df.iloc[-1]
    price = last["Close"]
    sma50, sma200 = last["SMA50"], last["SMA200"]
    s50, s200 = last["SMA50_slope20"], last["SMA200_slope20"]
    if pd.isna(sma50) or pd.isna(sma200):
        return {"stage": "Unknown", "sma50_color": "unknown", "sma200_color": "unknown",
                "sma50_slope": 0, "sma200_slope": 0}
    c50 = "green (rising)" if s50 > 0 else "red (declining)"
    c200 = "green (rising)" if s200 > 0 else "red (declining)"
    if price > sma50 > sma200 and s50 > 0 and s200 > 0:
        stage = "Stage 2 — Advancing"
    elif price < sma50 and price < sma200 and s50 < 0 and s200 < 0:
        stage = "Stage 4 — Declining"
    elif abs(s50) < price * 0.01:
        stage = "Stage 1 — Basing"
    else:
        stage = "Stage 3 — Topping / Transitioning"
    return {"stage": stage, "sma50_color": c50, "sma200_color": c200,
            "sma50_slope": s50, "sma200_slope": s200}


# ============================================================================
# VCP
# ============================================================================

def vcp_scorecard(df):
    last = df.iloc[-1]
    price = last["Close"]
    lb2y = df.tail(min(504, len(df)))
    base_low = float(lb2y["Low"].min())
    recent_high = float(lb2y["High"].max())
    uptrend_pct = (recent_high / base_low - 1) * 100 if base_low > 0 else 0

    lb1y = df.tail(min(252, len(df)))
    rh = float(lb1y["High"].max())
    hi_idx = lb1y["High"].idxmax()
    post = lb1y.loc[hi_idx:]
    rl = float(post["Low"].min()) if not post.empty else price
    depth_pct = (1 - rl / rh) * 100 if rh > 0 else 0
    try:
        weeks = (df.index[-1] - hi_idx).days / 7
    except Exception:
        weeks = 0
    recent_vol = df["Volume"].tail(10).mean()
    prior_vol = df["Volume"].tail(50).head(40).mean() if len(df) >= 50 else recent_vol

    checks = {
        "uptrend": uptrend_pct >= 25,
        "depth": depth_pct < 35,
        "duration": 5 <= weeks <= 65,
        "dry_up": recent_vol < prior_vol if prior_vol > 0 else False,
        "rvol": last["RVol"] >= 1.0 if not pd.isna(last["RVol"]) else False,
        "sma50_up": last["SMA50_slope20"] > 0,
        "sma200_up": last["SMA200_slope20"] > 0,
        "above_50": price > last["SMA50"] if not pd.isna(last["SMA50"]) else False,
    }
    p = sum(1 for v in checks.values() if v)
    t = len(checks)
    pct = p / t * 100
    if pct >= 85:   grade = "A+"
    elif pct >= 70: grade = "B"
    elif pct >= 50: grade = "C (mixed)"
    else:           grade = "Not Actionable"
    return {"checks": checks, "pass": p, "total": t, "grade": grade,
            "base_depth_pct": depth_pct, "base_weeks": weeks}


# ============================================================================
# KELL
# ============================================================================

def kell_cycle(df):
    last = df.iloc[-1]
    price = last["Close"]
    ema8, ema21 = last["EMA8"], last["EMA21"]
    atr_mult = last["ATRMult"] if not pd.isna(last["ATRMult"]) else 0
    last5 = df.tail(5)
    up5 = int((last5["Close"] > last5["Open"]).sum())
    rvol5 = float(last5["RVol"].mean()) if not last5["RVol"].isna().all() else 0
    above8 = price > ema8 if not pd.isna(ema8) else False
    above21 = price > ema21 if not pd.isna(ema21) else False
    cross_up = ema8 > ema21 if not (pd.isna(ema8) or pd.isna(ema21)) else False
    pct20 = (df["Close"].iloc[-1] / df["Close"].iloc[-20] - 1) * 100 if len(df) >= 20 else 0
    if atr_mult > 7 and rvol5 > 1.3:
        p = "Exhaustion Extension"
    elif atr_mult < -3 and up5 < 2:
        p = "Reversal Extension (in progress)"
    elif not above21 and above8 and up5 >= 3:
        p = "EMA Crossback (attempting)"
    elif above8 and above21 and cross_up and rvol5 > 1.0:
        p = "Base n Break (potential)"
    elif pct20 < -10 and up5 >= 3 and rvol5 > 1.2:
        p = "Wedge Pop (early)"
    else:
        p = "Neutral / Consolidation"
    return {"pattern": p, "atr_mult": atr_mult, "last5_up": up5, "last5_rvol": rvol5}


# ============================================================================
# FIB
# ============================================================================

def fib_retracement(df):
    lb = df.tail(min(504, len(df)))
    hi_val = float(lb["High"].max())
    hi_idx = lb["High"].idxmax()
    pre = lb.loc[:hi_idx]
    lo_val = float(pre["Low"].min()) if len(pre) > 20 else float(lb["Low"].min())
    rng = hi_val - lo_val
    cur = float(df["Close"].iloc[-1])
    retrace = (hi_val - cur) / rng * 100 if rng > 0 else 0
    return {"high": hi_val, "low": lo_val, "range": rng, "current": cur,
            "retrace_pct": retrace}


# ============================================================================
# WEEKLY VOLUME PROFILE
# ============================================================================

def weekly_profile(df, num_weeks=12):
    weekly = df.resample("W").agg({
        "High": "max", "Low": "min", "Open": "first", "Close": "last", "Volume": "sum",
    }).tail(num_weeks).dropna()
    out = []
    for wk_end in weekly.index:
        wk_start = wk_end - pd.Timedelta(days=7)
        wk_bars = df[(df.index > wk_start) & (df.index <= wk_end)]
        if len(wk_bars) == 0:
            continue
        lo, hi = wk_bars["Low"].min(), wk_bars["High"].max()
        if hi <= lo:
            continue
        bw = (hi - lo) / 20
        buckets = {}
        for _, bar in wk_bars.iterrows():
            typ = (bar["High"] + bar["Low"] + bar["Close"]) / 3
            b = min(max(int((typ - lo) / bw) if bw > 0 else 0, 0), 19)
            buckets[b] = buckets.get(b, 0) + bar["Volume"]
        if not buckets:
            continue
        poc_b = max(buckets, key=buckets.get)
        poc = lo + (poc_b + 0.5) * bw
        tot = sum(buckets.values())
        vas = {poc_b}
        vvol = buckets[poc_b]
        up, dn = poc_b + 1, poc_b - 1
        while vvol < tot * 0.7 and (up <= 19 or dn >= 0):
            u = buckets.get(up, 0) if up <= 19 else -1
            d = buckets.get(dn, 0) if dn >= 0 else -1
            if u >= d and up <= 19:
                vas.add(up); vvol += u; up += 1
            elif dn >= 0:
                vas.add(dn); vvol += d; dn -= 1
            else:
                break
        vah = lo + (max(vas) + 1) * bw
        val = lo + min(vas) * bw
        out.append({"week_end": wk_end.date(), "high": float(hi), "low": float(lo),
                    "poc": float(poc), "vah": float(vah), "val": float(val)})
    return out


def vah_ladder_analysis(profile):
    if len(profile) < 3:
        return "Insufficient weekly data"
    vahs = [w["vah"] for w in profile[-4:]]
    vals = [w["val"] for w in profile[-4:]]
    rising = len(vahs) >= 3 and vahs[-1] > vahs[-2] > vahs[-3]
    falling = len(vahs) >= 3 and vahs[-1] < vahs[-2] < vahs[-3]
    breakout = len(vals) >= 2 and vals[-1] > vahs[-2]
    breakdown = len(vahs) >= 2 and vahs[-1] < vals[-2]
    if breakout:
        return "VALUE BREAKOUT — current VAL above prior week's VAH"
    if breakdown:
        return "VALUE BREAKDOWN — current VAH below prior week's VAL"
    if rising:
        return "VAH ladder rising — value migrating up"
    if falling:
        return "VAH ladder falling — value migrating down"
    return "VAH sideways"


# ============================================================================
# EPISODIC PIVOT
# ============================================================================

def episodic_pivot_check(df, news):
    recent = df.tail(20)
    candidates = recent[(recent["RVol"] > 1.5) & (recent["PctChg"].abs() > 5)]
    if candidates.empty:
        return {"detected": False, "events": [], "summary": "No recent EP candle"}
    events = []
    for idx, bar in candidates.iterrows():
        bd = idx.date()
        matched = [n for n in news if abs((n["date"].date() - bd).days) <= 3]
        events.append({"date": bd, "pct_chg": bar["PctChg"], "rvol": bar["RVol"],
                       "news": matched})
    any_news = any(e["news"] for e in events)
    return {"detected": True, "events": events,
            "summary": "EP candle with news catalyst" if any_news else
                       f"{len(events)} anomalous bar(s) but no news match"}


# ============================================================================
# CLAUDE VISION
# ============================================================================

def analyze_chart_with_vision(image_path, ticker, computed_summary, api_key=None,
                              model="claude-sonnet-4-5"):
    import base64

    # Diagnostic banner so the user clearly sees what's happening
    print()
    print("=" * 72)
    print("  CHART VISION ANALYSIS")
    print("=" * 72)

    if not api_key:
        api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("  [SKIPPED] No Anthropic API key found.")
        print("  To enable chart vision:")
        print("    1. Get a key from https://console.anthropic.com/")
        print("    2. Run this in your terminal:")
        print("         set ANTHROPIC_API_KEY=sk-ant-...")
        print("    3. Re-run this script")
        print("=" * 72)
        return None

    if not os.path.isfile(image_path):
        print(f"  [SKIPPED] Chart file not found: {image_path}")
        print("=" * 72)
        return None

    # Check anthropic SDK availability
    try:
        import anthropic
        sdk_available = True
    except ImportError:
        sdk_available = False
        print("  [WARN] 'anthropic' package not installed — using raw HTTP fallback")
        print("         Install with: pip install anthropic")

    print(f"  [OK] API key detected ({'sk-ant-...' + api_key[-4:]})")
    print(f"  [OK] Chart file: {image_path}")
    print(f"  [OK] SDK: {'anthropic' if sdk_available else 'urllib fallback'}")
    print(f"  [OK] Model: {model}")
    print(f"  Sending request to Claude...")
    print("=" * 72)
    with open(image_path, "rb") as f:
        data = base64.standard_b64encode(f.read()).decode("utf-8")
    ext = os.path.splitext(image_path)[1].lower()
    mt = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
          ".gif": "image/gif", ".webp": "image/webp"}.get(ext, "image/png")

    prompt = f"""Here is a computed technical read of {ticker}:

{computed_summary}

Now analyze the attached TradingView chart. Keep output to ONE concise paragraph
describing what you see visually that confirms or contradicts the computed read,
plus any custom indicators visible (ATR Multiple bubbles, weekly POC boxes,
virgin POCs, red/green dotted swing lines, Pocket Pivot volume colors).
Name specific reclaim price levels (Tier 1/2/3). Do NOT use checkbox tables,
just natural-language commentary. Maximum 200 words."""

    result = None
    try:
        import anthropic
        client = anthropic.Anthropic(api_key=api_key)
        msg = client.messages.create(
            model=model, max_tokens=1500,
            messages=[{"role": "user", "content": [
                {"type": "image", "source": {"type": "base64", "media_type": mt, "data": data}},
                {"type": "text", "text": prompt},
            ]}],
        )
        result = msg.content[0].text
    except ImportError:
        # anthropic package missing, use raw urllib
        import urllib.request, urllib.error, json as _json
        body = {"model": model, "max_tokens": 1500, "messages": [{"role": "user", "content": [
            {"type": "image", "source": {"type": "base64", "media_type": mt, "data": data}},
            {"type": "text", "text": prompt},
        ]}]}
        req = urllib.request.Request(
            "https://api.anthropic.com/v1/messages",
            data=_json.dumps(body).encode("utf-8"),
            headers={"x-api-key": api_key, "anthropic-version": "2023-06-01",
                     "content-type": "application/json"})
        try:
            with urllib.request.urlopen(req, timeout=90) as resp:
                result = _json.loads(resp.read())["content"][0]["text"]
        except urllib.error.HTTPError as e:
            body = e.read().decode("utf-8", errors="replace")
            print(f"[vision] API HTTP {e.code}: {body[:300]}")
        except Exception as e:
            print(f"[vision] API failed: {e}")
    except anthropic.APIStatusError as e:
        print(f"[vision] API status error: {e}")
    except Exception as e:
        print(f"[vision] API failed: {type(e).__name__}: {e}")

    if result:
        print(f"[vision] Got {len(result)} chars from Claude — visual read included in brief.")
    else:
        print(f"[vision] FAILED — visual read will be missing from brief.")
    return result


# ============================================================================
# SIGNAL AGGREGATOR + ACTION LOGIC
# ============================================================================

def classify_signals(df, stage_info, vcp, kell, fib, vah_analysis, structure,
                     prior_hi, prior_lo, reclaim, ep):
    """Return (bullish_list, bearish_list, action_paragraph, action_tag)."""
    last = df.iloc[-1]
    bull, bear = [], []

    stage = stage_info["stage"]
    if stage.startswith("Stage 2"):
        bull.append("price is above a rising 50MA and 200MA — a textbook Stage 2 uptrend")
    elif stage.startswith("Stage 1"):
        if stage_info.get("sma50_slope", 0) > 0:
            bull.append("the 50MA has started curling up from a flat base — early Stage 2 trigger setting up")
        else:
            bear.append("the stock is stuck in a flat Stage 1 base with no directional edge yet")
    elif stage.startswith("Stage 3"):
        bear.append("Stage 3 topping behavior — the 50MA is rolling over and distribution is visible")
    elif stage.startswith("Stage 4"):
        bear.append("Stage 4 decline — price is below a falling 50MA and 200MA, still trending down")

    if "red" in stage_info.get("sma200_color", ""):
        bear.append("the 200MA is red and sloping down, acting as dynamic overhead resistance on any rally")
    elif "green" in stage_info.get("sma200_color", ""):
        bull.append("the 200MA is rising, confirming the long-term trend is supportive")

    grade = vcp.get("grade", "")
    if grade.startswith("A"):
        bull.append("the base passes most Minervini VCP criteria — institutional-grade setup")
    elif "Not Actionable" in grade:
        bear.append("the VCP scorecard fails — base depth, duration, or structure are not yet tradeable")
    elif grade.startswith("C"):
        bear.append("the VCP is mixed — quality is there but with meaningful gaps")

    if vcp.get("base_depth_pct", 0) > 40:
        bear.append(f"the base is {vcp['base_depth_pct']:.0f}% deep, carrying heavy overhead supply")

    if "Uptrend" in structure:
        bull.append("market structure remains intact with higher highs and higher lows")
    elif "Reclaimed" in structure:
        bull.append(f"price has reclaimed the prior swing high around ${prior_hi:.2f}, breaking the downtrend structure")
    elif "Downtrend" in structure:
        bear.append("the downtrend structure of lower highs and lower lows has not yet been invalidated")

    kp = kell.get("pattern", "")
    if "Base n Break" in kp:
        bull.append("an Oliver Kell Base-n-Break is triggering — above both EMAs on expanding volume")
    elif "EMA Crossback" in kp:
        bull.append("price is reclaiming the 8/21 EMAs from below — a Kell EMA Crossback swing trigger")
    elif "Wedge Pop" in kp:
        bull.append("an early Wedge Pop is forming — sharp prior decline followed by expanding green candles")
    elif "Reversal Extension" in kp:
        bear.append("the Kell Reversal Extension is still playing out — no bottom confirmation yet")
    elif "Exhaustion" in kp:
        bear.append("a Kell Exhaustion Extension is in play — parabolic blow-off risk, not a long")

    am = kell.get("atr_mult", 0)
    if abs(am) < 3 and stage.startswith(("Stage 1", "Stage 2")):
        bull.append("the stock is not extended from its 50MA — plenty of room to run before any blow-off risk")
    elif am >= 7:
        bear.append("price is extended into the ATR blow-off zone — do not chase")

    if fib:
        rp = fib.get("retrace_pct", 0)
        if rp < 38.2:
            bull.append("only a shallow Fibonacci pullback — signature of a strong stock")
        elif rp < 50:
            bull.append("the Fibonacci retracement is in Minervini's preferred healthy-pullback zone")
        elif rp < 61.8:
            bull.append("the pullback bottomed near the 61.8% golden pocket — classic deep-but-acceptable retrace")
        elif rp < 78.6:
            bear.append("the stock has retraced beyond the 61.8% golden pocket — stressed retracement")
        else:
            bear.append("nearly a full Fibonacci round-trip — a stage reset is needed before long setups")

    if "BREAKOUT" in vah_analysis:
        bull.append("the weekly VAL just broke above the prior week's VAH — a clean value breakout")
    elif "BREAKDOWN" in vah_analysis:
        bear.append("the weekly VAH broke below the prior week's VAL — ongoing value breakdown")
    elif "rising" in vah_analysis.lower():
        bull.append("the weekly VAH ladder is climbing — value migrating up (Wyckoff markup)")
    elif "falling" in vah_analysis.lower():
        bear.append("the weekly VAH ladder is falling — value migrating down")

    recent20 = df.tail(20)
    colors = recent20["VolColor"].value_counts().to_dict()
    blue, green, red, orange, gray = (colors.get(c, 0) for c in
                                      ("blue", "green", "red", "orange", "gray"))
    if blue >= 2 and orange >= 1:
        bull.append("an orange-dry → blue Pocket Pivot sequence has printed — classic accumulation signature")
    elif blue >= 2:
        bull.append("multiple Pocket Pivot volume bars recently — institutions are stepping in")
    if red > green + blue + 2:
        bear.append("red distribution bars dominate the last 20 sessions — institutions selling")
    if gray >= 14:
        bear.append("most recent volume is noise-grade gray bars — no institutional commitment in either direction")

    hvq_recent = df.tail(63)[df.tail(63)["HVQ"]]
    hvy_recent = df.tail(252)[df.tail(252)["HVY"]]
    if not hvy_recent.empty:
        bull.append(f"Highest Volume of Year bar printed on {hvy_recent.index[-1].date()} — strongest institutional buying signal available")
    elif not hvq_recent.empty:
        bull.append(f"Highest Volume of Quarter bar printed on {hvq_recent.index[-1].date()}")

    if ep.get("detected") and any(e.get("news") for e in ep.get("events", [])):
        bull.append("a recent large-volume bar coincides with a news catalyst — Episodic Pivot in play")

    # Decision
    score = len(bull) - len(bear)
    t1 = reclaim["tier1"]["price"]
    t2 = reclaim["tier2"]["price"]
    t3 = reclaim["tier3"]["price"]
    stop = prior_lo if prior_lo else float(df["Low"].tail(10).min())
    sma50 = last["SMA50"]

    if stage.startswith("Stage 4") or ("Not Actionable" in grade and score < 1):
        tag = "NO TRADE"
        action = (
            f"No hurry. This is a broken chart with overhead supply and no clean accumulation signature yet. "
            f"Put it on the watchlist and revisit only after a decisive close above ${t3:.2f} on expanding blue/green volume. "
            f"Any long attempt before then is fighting the tape — invalidation sits below ${stop:.2f}."
        )
    elif stage.startswith("Stage 2") and score >= 3:
        tag = "BUY on pullback"
        action = (
            f"Actionable long. The preferred entry is a pullback to the rising 50MA near ${sma50:.2f}, "
            f"with a tight stop below ${stop:.2f}. If there is no pullback and price breaks out through Tier 1 ${t1:.2f} "
            f"on above-average blue volume, scale in with a partial. Reserve full position size for a decisive reclaim "
            f"of Tier 3 ${t3:.2f} — the structural breakout trigger."
        )
    elif score >= 2 and (stage.startswith("Stage 1") or "Reclaimed" in structure):
        tag = "WATCH — scale in on reclaim"
        action = (
            f"Constructive but not confirmed. Do not buy here; scale in only on reclaim tiers. "
            f"Toe in at Tier 1 ${t1:.2f} (20% size), add at Tier 2 ${t2:.2f} (to 50%), and take full size only above Tier 3 ${t3:.2f}. "
            f"Stop below ${stop:.2f}. Consolidation can drag on for weeks — patience until the reclaim prints on volume."
        )
    elif score >= 1:
        tag = "WATCH"
        action = (
            f"Mixed with a slight bullish lean. No hurry — wait for a decisive reclaim of Tier 2 ${t2:.2f} on strong volume "
            f"before committing capital. Stop reference ${stop:.2f}. A failed reclaim back below Tier 1 ${t1:.2f} "
            f"means it needs more time."
        )
    else:
        tag = "NO TRADE"
        action = (
            f"Not a buy here. Bearish signals outnumber bullish and there's no reclaim trigger yet. "
            f"Wait for a clean close above Tier 3 ${t3:.2f} with volume expansion before reconsidering. "
            f"Invalidation below ${stop:.2f}."
        )

    return bull, bear, action, tag


def build_paragraph(items, opener):
    """Combine signal list into a readable paragraph."""
    if not items:
        return f"{opener} nothing meaningful at this point."
    if len(items) == 1:
        return f"{opener} {items[0]}."
    # Join with commas; last one with "and"
    body = "; ".join(items[:-1]) + f"; and {items[-1]}"
    return f"{opener} {body}."


# ============================================================================
# DOCX REPORT
# ============================================================================

def _set_cell_bg(cell, rgb_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), rgb_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def generate_report(ticker, df, bull, bear, action, tag, reclaim,
                    prior_lo, chart_path=None, vision_md=None,
                    output_path=None):
    doc = Document()

    # Title
    h = doc.add_heading(f"{ticker} — Trade Brief", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle / verdict tag
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(tag)
    r.bold = True
    r.font.size = Pt(16)
    color_map = {
        "BUY on pullback": RGBColor(0x22, 0x8B, 0x22),
        "WATCH — scale in on reclaim": RGBColor(0xE0, 0x80, 0x00),
        "WATCH": RGBColor(0xE0, 0x80, 0x00),
        "NO TRADE": RGBColor(0xC6, 0x28, 0x28),
    }
    r.font.color.rgb = color_map.get(tag, RGBColor(0x33, 0x33, 0x33))

    # Date + price
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"${df.iloc[-1]['Close']:.2f}  ·  {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_paragraph()

    # ── Paragraph 1: Bullish picture ──
    _heading(doc, "Bullish picture", level=2)
    p = doc.add_paragraph()
    p.add_run(build_paragraph(bull, "On the bullish side,")).font.size = Pt(11)

    # ── Paragraph 2: Bearish + Action ──
    _heading(doc, "Bearish picture & action", level=2)
    p = doc.add_paragraph()
    p.add_run(build_paragraph(bear, "Against that,")).font.size = Pt(11)
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(action)
    run.font.size = Pt(11)
    run.bold = True

    # ── Reclaim levels (compact box) ──
    doc.add_paragraph()
    _heading(doc, "Reclaim levels for action", level=2)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Light Grid Accent 1"
    headers = ["Tier", "Price", "Size on trigger"]
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
        _set_cell_bg(tbl.rows[0].cells[i], "1F3A5F")
        for r in tbl.rows[0].cells[i].paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for label, key, size in [
        ("Tier 1 — tactical", "tier1", "20%"),
        ("Tier 2 — intermediate", "tier2", "50%"),
        ("Tier 3 — structural", "tier3", "100% (full)"),
    ]:
        row = tbl.add_row().cells
        row[0].text = label
        row[1].text = f"${reclaim[key]['price']:.2f}"
        row[2].text = size

    if prior_lo:
        p = doc.add_paragraph()
        r = p.add_run(f"Invalidation: close below ${prior_lo:.2f} on distribution volume.")
        r.italic = True
        r.font.size = Pt(10)

    # ── Chart + Vision (if provided) ──
    if chart_path and os.path.isfile(chart_path):
        doc.add_paragraph()
        _heading(doc, "Chart", level=2)
        try:
            doc.add_picture(chart_path, width=Inches(6.5))
        except Exception as e:
            doc.add_paragraph(f"(Could not embed image: {e})")

    if vision_md:
        doc.add_paragraph()
        _heading(doc, "Visual read", level=2)
        p = doc.add_paragraph()
        p.add_run(vision_md).font.size = Pt(10)

    # Disclaimer
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Not financial advice. Framework: Minervini VCP + Oliver Kell + weekly POC + Pocket Pivot Volume + Episodic Pivot + ATR Multiple.")
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(output_path)
    return output_path


# ============================================================================
# MAIN
# ============================================================================

def _ansi_color_for_tag(tag):
    """Return ANSI color code for a verdict tag."""
    if tag.startswith("BUY"):
        return "\033[92m"   # bright green
    if tag.startswith("WATCH"):
        return "\033[93m"   # bright yellow
    if tag == "NO TRADE":
        return "\033[91m"   # bright red
    return "\033[0m"


def print_brief(ticker, df, bull, bear, action, tag, reclaim, prior_lo,
                vision_md=None):
    """Print the concise trade brief directly to the terminal."""
    C = _ansi_color_for_tag(tag)
    BOLD = "\033[1m"
    DIM = "\033[2m"
    UND = "\033[4m"
    R = "\033[0m"   # reset
    CYAN = "\033[96m"

    width = 72
    last = df.iloc[-1]["Close"]

    print()
    print(BOLD + "=" * width + R)
    print(BOLD + f"  {ticker} — TRADE BRIEF".center(width) + R)
    print(BOLD + "=" * width + R)
    print(f"  Verdict:  {C}{BOLD}{tag}{R}")
    print(f"  Price:    ${last:.2f}")
    print(f"  Time:     {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    print("-" * width)
    print()

    # Bullish paragraph
    print(f"{UND}{BOLD}BULLISH PICTURE{R}")
    bull_text = build_paragraph(bull, "On the bullish side,")
    print(_wrap(bull_text, width))
    print()

    # Bearish + action
    print(f"{UND}{BOLD}BEARISH PICTURE & ACTION{R}")
    bear_text = build_paragraph(bear, "Against that,")
    print(_wrap(bear_text, width))
    print()
    print(f"{BOLD}{C}Action:{R}")
    print(_wrap(action, width))
    print()

    # Reclaim table (simple text)
    print(f"{UND}{BOLD}RECLAIM LEVELS FOR ACTION{R}")
    print(f"  Tier 1 (tactical)      ${reclaim['tier1']['price']:>8.2f}   →  20% size")
    print(f"  Tier 2 (intermediate)  ${reclaim['tier2']['price']:>8.2f}   →  50% size")
    print(f"  Tier 3 (structural)    ${reclaim['tier3']['price']:>8.2f}   → 100% (full)")
    if prior_lo:
        print(f"  {DIM}Invalidation: close below ${prior_lo:.2f} on distribution volume{R}")
    print()

    # Visual read (if chart was given)
    if vision_md:
        print(f"{UND}{BOLD}VISUAL READ (Claude vision){R}")
        print(_wrap(vision_md, width))
        print()

    print(BOLD + "=" * width + R)
    print()


def _wrap(text, width=72, indent=""):
    """Simple word-wrap for terminal output."""
    import textwrap
    return "\n".join(textwrap.fill(line, width=width, initial_indent=indent,
                                    subsequent_indent=indent)
                     for line in text.split("\n") if line.strip() or True)


def run(ticker, output_path=None, period="3y", source="tws",
        tws_host="127.0.0.1", tws_port=7497, tws_client=20,
        chart_path=None, api_key=None, save_docx=False):
    ticker = ticker.upper().strip()

    print(f"\n=== Analyzing {ticker} ===\n")
    df, ticker_obj = fetch_data(ticker, period, source, tws_host, tws_port, tws_client)
    df = add_indicators(df)
    df = classify_volume_colors(df)
    print("[indicators] Computed MAs, ATR, PPV colors, HVQ/HVY")

    sw_hi, sw_lo = find_swings(df)
    structure, prior_hi, prior_lo = determine_structure(df, sw_hi, sw_lo)
    reclaim = determine_reclaim_levels(df, sw_hi)
    stage_info = determine_stage(df)
    vcp = vcp_scorecard(df)
    kell = kell_cycle(df)
    fib = fib_retracement(df)
    profile = weekly_profile(df)
    vah = vah_ladder_analysis(profile)
    news = fetch_news(ticker_obj, 45) if ticker_obj else []
    ep = episodic_pivot_check(df, news)
    print(f"[analysis] Stage={stage_info['stage']} | VCP={vcp['grade']} | Kell={kell['pattern']}")

    bull, bear, action, tag = classify_signals(
        df, stage_info, vcp, kell, fib, vah, structure,
        prior_hi, prior_lo, reclaim, ep,
    )
    print(f"[verdict] {tag} | {len(bull)} bullish, {len(bear)} bearish signals")

    # Optional Claude vision on chart
    vision_md = None
    if chart_path:
        lines = [
            f"Verdict tag: {tag}",
            f"Bullish: {' | '.join(bull) if bull else 'none'}",
            f"Bearish: {' | '.join(bear) if bear else 'none'}",
            f"Reclaim Tiers: T1 ${reclaim['tier1']['price']:.2f}, "
            f"T2 ${reclaim['tier2']['price']:.2f}, T3 ${reclaim['tier3']['price']:.2f}",
        ]
        vision_md = analyze_chart_with_vision(
            chart_path, ticker, "\n".join(lines), api_key=api_key,
        )

    # Print to terminal (primary output)
    print_brief(ticker, df, bull, bear, action, tag, reclaim, prior_lo, vision_md)

    # Optional DOCX save
    if save_docx:
        if output_path is None:
            output_path = f"{ticker}_analysis_{datetime.now().strftime('%Y%m%d')}.docx"
        generate_report(ticker, df, bull, bear, action, tag, reclaim,
                        prior_lo, chart_path, vision_md, output_path)
        print(f"[done] DOCX saved -> {output_path}\n")

    return tag


def interactive_loop():
    """Interactive prompt mode — runs when script is double-clicked (no CLI args)."""
    while True:
        print()
        print("=" * 72)
        print("  STOCK ANALYZER - Minervini + Kell + POC + PPV + EP + ATR".center(72))
        print("=" * 72)
        print()
        try:
            ticker = input("Enter ticker symbol (or blank to quit): ").strip()
        except (EOFError, KeyboardInterrupt):
            break
        if not ticker:
            break

        try:
            chart_choice = input("Attach a chart image? [y/N]: ").strip().lower()
        except (EOFError, KeyboardInterrupt):
            break

        chart_path = None
        if chart_choice == "y":
            try:
                chart_path = input("Drag-drop the chart file here and press ENTER: ").strip()
                # Strip quotes from drag-drop
                chart_path = chart_path.strip('"').strip("'")
            except (EOFError, KeyboardInterrupt):
                chart_path = None

        try:
            run(ticker, chart_path=chart_path)
        except Exception as e:
            print()
            print("=" * 72)
            print(f"[ERROR] {type(e).__name__}: {e}")
            print("=" * 72)
            import traceback
            traceback.print_exc()

        print()
        try:
            again = input("Press ENTER to analyze another ticker, or type Q to quit: ").strip().lower()
        except (EOFError, KeyboardInterrupt):
            break
        if again == "q":
            break
        # Clear screen for next run (Windows or Unix)
        os.system("cls" if os.name == "nt" else "clear")

    print("\nGoodbye.")


def cli():
    # If no arguments were given, fall into interactive mode (useful when
    # the script is double-clicked from Explorer).
    if len(sys.argv) == 1:
        try:
            interactive_loop()
        except KeyboardInterrupt:
            pass
        return

    parser = argparse.ArgumentParser(description="Stock Technical Analysis -> DOCX brief")
    parser.add_argument("ticker", help="Stock ticker (e.g. AAPL, HOOD, SE)")
    parser.add_argument("--output", "-o", default=None, help="Output DOCX path")
    parser.add_argument("--period", "-p", default="3y",
                        help="Data period: 1y / 2y / 3y (default) / 5y / 10y / max")
    parser.add_argument("--yfinance", action="store_true", help="Force yfinance")
    parser.add_argument("--tws-only", action="store_true", help="Require TWS, no fallback")
    parser.add_argument("--tws-host", default="127.0.0.1")
    parser.add_argument("--tws-port", type=int, default=7497)
    parser.add_argument("--tws-client", type=int, default=20)
    parser.add_argument("--chart", "-c", default=None,
                        help="Chart image path (triggers Claude vision)")
    parser.add_argument("--api-key", default=None,
                        help="Anthropic API key (or set ANTHROPIC_API_KEY env var)")
    parser.add_argument("--docx", action="store_true",
                        help="Also save the brief to a DOCX file")
    args = parser.parse_args()

    source = "yfinance" if args.yfinance else ("tws_only" if args.tws_only else "tws")

    exit_code = 0
    try:
        run(args.ticker, output_path=args.output, period=args.period,
            source=source, tws_host=args.tws_host, tws_port=args.tws_port,
            tws_client=args.tws_client, chart_path=args.chart,
            api_key=args.api_key, save_docx=args.docx)
    except KeyboardInterrupt:
        print("\n[INTERRUPTED]")
        exit_code = 130
    except Exception as e:
        print("\n" + "=" * 70)
        print(f"[ERROR] {type(e).__name__}: {e}")
        print("=" * 70)
        import traceback
        traceback.print_exc()
        exit_code = 1

    # Skip pause when launched from run_stock_analyzer.bat (it has its own loop)
    skip_pause = os.environ.get("STOCK_ANALYZER_NO_PAUSE") == "1"
    if not skip_pause and sys.stdin.isatty() and sys.stdout.isatty():
        try:
            input("\nPress ENTER to close this window...")
        except (EOFError, KeyboardInterrupt):
            pass
    sys.exit(exit_code)


if __name__ == "__main__":
    cli()
